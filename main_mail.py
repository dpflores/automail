import time                             # importing time
from docx import Document               # to edit word
import os                               # for files input
from datetime import date               # Import the date
import smtplib                          # for mailing
from email.message import EmailMessage  # for message mailing

# for windows
try:
    from docx2pdf import convert        # to convert to pdf
except:
    pass

EMAIL_ADDRESS = os.environ.get('EMAIL_USER')    #delpiero22
EMAIL_PASSWORD = os.environ.get('EMAIL_PASS')  # Email Password for python
RECEIVER = os.environ.get('RECEIVER')




TEMPLATE_WORD = 'plantilla.docx'
OUTPUT_NAME = 'recibo_de_conformidad'
OUTPUT_NAME_WORD = OUTPUT_NAME + '.docx'
OUTPUT_NAME_PDF = OUTPUT_NAME + '.pdf'


# MAILING VARIABLES
msg = EmailMessage()
msg['Subject'] = 'Recibo de conformidad'
msg['From'] = EMAIL_ADDRESS
msg['To'] = RECEIVER
msg.set_content('Hola, adjunto el recibo de conformidad del correspondiente mes')

# CONVERSION VARIABLES
# creating the date object of today's date
todays_date = date.today()
months_list = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio","agosto", "setiembre", "octubre", "noviembre","diciembre"]

VARIABLES = {
        "${DAY_NUMBER}": str(todays_date.day),
        "${MONTH_NUMBER}": str(todays_date.month),
        "${YEAR_NUMBER}": str(todays_date.year),
        "${MONTH_WORD}": months_list[todays_date.month - 1]
    }

# FUNCTIONS
def edit_word(template_file_path, output_file_path, variables):
    template_document = Document(template_file_path)
    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)
    template_document.save(output_file_path)

def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key,value)

def send_mail(file, email_adress, email_password, receiver):
    with open(file, 'rb') as f:
        file_data = f.read()
        file_name = f.name
    msg.add_attachment(file_data, maintype='application',
                        subtype='octet-stream', filename=file_name)
    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
        smtp.login(email_adress, email_password)
        smtp.send_message(msg)


def main():
    print("Envio de recibo de conformidad mensual al BCP")


    # Word template edition
    print("Editing template...")
    edit_word(TEMPLATE_WORD, OUTPUT_NAME_WORD, VARIABLES)

    # Word conversion to pdf
    print("Converting word to pdf...")

    # doc = Document(OUTPUT_NAME_WORD)
    # doc.save(OUTPUT_NAME_PDF)

    command = "lowriter --convert-to pdf "+ OUTPUT_NAME_WORD

    #linux
    try:
        os.system(command)

    #windows
    except:
        convert(OUTPUT_NAME_WORD)
        convert(OUTPUT_NAME_WORD, OUTPUT_NAME_PDF)


    #Sending file by email
    print("Sending file by email...")
    send_mail(OUTPUT_NAME_PDF, EMAIL_ADDRESS, EMAIL_PASSWORD, RECEIVER)
    print("Email sended to: %s" %(RECEIVER))
    time.sleep(3)

if __name__ == '__main__':
    main()
