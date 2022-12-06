# importing to edit pdf
from docx import Document

# to search the teamplate of word
import os

# importing date class from datetime module
from datetime import date

# convert word to pdf
from docx2pdf import convert

OUTPUT_NAME = 'recibo_de_conformidad'
OUTPUT_NAME_WORD = OUTPUT_NAME + '.docx'
OUTPUT_NAME_PDF = OUTPUT_NAME + '.pdf'
TEMPLATE_WORD = 'plantilla.docx'

# creating the date object of today's date
todays_date = date.today()

months_list = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio","agosto", "setiembre", "octubre", "noviembre","diciembre"]

def main():
    template_file_path = TEMPLATE_WORD
    output_file_path = OUTPUT_NAME_WORD

    variables = {
        "${DAY_NUMBER}": str(todays_date.day),
        "${MONTH_NUMBER}": str(todays_date.month),
        "${YEAR_NUMBER}": str(todays_date.year),
        "${MONTH_WORD}": months_list[todays_date.month - 1]
    }

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        '''for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)'''

    template_document.save(output_file_path)


    # Conversion to pdf
    convert(OUTPUT_NAME_WORD, OUTPUT_NAME_PDF)

def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key,value)


if __name__ == '__main__':
    main()
